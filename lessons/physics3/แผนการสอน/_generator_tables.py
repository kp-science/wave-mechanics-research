#!/usr/bin/env python3
"""
Generator สำหรับ 3 ตารางเอกสารประกอบหลักสูตร (ฟิสิกส์ 3 · ว 30203 · ม.5)
- 04 ตารางการวิเคราะห์มาตรฐาน ตัวชี้วัด
- 05 ตารางการออกแบบหน่วยการเรียนรู้
- 08 ตารางแผนการจัดการเรียนรู้

Output: lessons/physics3/แผนการสอน/เอกสารประกอบ/*.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pythainlp.tokenize import word_tokenize

ZWSP = "​"
FONT = "TH SarabunPSK"

OUTDIR = Path("/Users/komanepapato/Documents/วิจัย/wave-mechanics-research/lessons/physics3/แผนการสอน/เอกสารประกอบ")

THAI_NUMS = "๐๑๒๓๔๕๖๗๘๙"
def thai_num(n):
    return "".join(THAI_NUMS[int(d)] if d.isdigit() else d for d in str(n))

def th_break(s):
    if not s: return ""
    try:
        toks = word_tokenize(s, engine="newmm", keep_whitespace=True)
        return ZWSP.join(toks)
    except Exception:
        return s

def set_font(run, size=14, bold=False, color=None, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfont = rpr.find(qn("w:rFonts"))
    if rfont is None:
        rfont = OxmlElement("w:rFonts"); rpr.append(rfont)
    for attr in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"):
        rfont.set(qn(attr), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(doc, text, *, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             color=None, space_before=0, space_after=2, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(th_break(text))
    set_font(r, size=size, bold=bold, color=color, italic=italic)
    return p

def shade(cell, hex_color):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(sh)

def set_cell(cell, text, *, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             color=None, shade_color=None, vertical=WD_ALIGN_VERTICAL.CENTER):
    cell.text = ""
    cell.vertical_alignment = vertical
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    # multi-line support
    lines = str(text).split("\n")
    for i, ln in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph(); p.alignment = align
            p.paragraph_format.space_after = Pt(1)
        r = p.add_run(th_break(ln))
        set_font(r, size=size, bold=bold, color=color)
    if shade_color:
        shade(cell, shade_color)

def set_landscape(doc):
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.5); sec.right_margin = Cm(1.5)

def add_signatures(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(20)
    for line in ["ลงชื่อ ..............................................  ครูผู้สอน",
                 "ลงชื่อ ..............................................  หัวหน้ากลุ่มสาระการเรียนรู้",
                 "ลงชื่อ ..............................................  รองผู้อำนวยการกลุ่มบริหารวิชาการ"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(line)
        set_font(r, size=14)

def add_header_block(doc):
    """หัวกระดาษ — กลุ่มสาระ / รายวิชา / หน่วยกิต / เวลา / ชั้น"""
    t = doc.add_table(rows=2, cols=3)
    t.autofit = False
    for col in t.columns:
        for cell in col.cells:
            cell._tc.tcPr.append(OxmlElement("w:tcBorders"))  # no borders
    # row 1
    set_cell(t.rows[0].cells[0], "กลุ่มสาระการเรียนรู้ วิทยาศาสตร์และเทคโนโลยี",
             size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(t.rows[0].cells[1], "รายวิชา ฟิสิกส์ 3 (ว 30203)",
             size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(t.rows[0].cells[2], "สาระเพิ่มเติม",
             size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
    # row 2
    set_cell(t.rows[1].cells[0], "จำนวน 2 หน่วยกิต",
             size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(t.rows[1].cells[1], "เวลา 4 คาบ/สัปดาห์",
             size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell(t.rows[1].cells[2], "ชั้นมัธยมศึกษาปีที่ 5",
             size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
    # remove borders
    tbl = t._tbl
    for cell in tbl.iter(qn("w:tc")):
        tcPr = cell.find(qn("w:tcPr"))
        if tcPr is not None:
            for tag in list(tcPr):
                if tag.tag == qn("w:tcBorders"):
                    tcPr.remove(tag)
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top","left","bottom","right"):
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"), "nil")
                tcBorders.append(b)
            tcPr.append(tcBorders)
    return t

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    set_font(r, size=18, bold=True)

# ─── เนื้อหาฟิสิกส์ 3 ────────────────────────────────────────
# ผลการเรียนรู้ ของรายวิชา ฟิสิกส์ 3 (ว 30203) · เพิ่มเติม · ม.5
LEARN_OUTCOMES = [
    {"code": "1", "text": "ทดลองและอธิบายการเกิดและสมบัติของคลื่นกล อัตราเร็วของคลื่น ความสัมพันธ์ระหว่างความถี่ ความยาวคลื่น และอัตราเร็วของคลื่นได้"},
    {"code": "2", "text": "ทดลองและอธิบายปรากฏการณ์ของคลื่น ได้แก่ การสะท้อน การหักเห การแทรกสอด การเลี้ยวเบน และคลื่นนิ่งได้"},
    {"code": "3", "text": "ทดลองและอธิบายการเคลื่อนที่แบบฮาร์มอนิกอย่างง่าย คาบ ความถี่ พลังงาน และการสั่นพ้องได้"},
    {"code": "4", "text": "อธิบายการเกิดและการเคลื่อนที่ของเสียง สมบัติของเสียง และปริมาณทางฟิสิกส์ที่เกี่ยวข้องกับเสียงได้"},
    {"code": "5", "text": "ทดลองและอธิบายปรากฏการณ์ของเสียง ได้แก่ การสั่นพ้อง บีต และปรากฏการณ์ดอปเพลอร์ได้"},
    {"code": "6", "text": "อธิบายระดับความเข้มเสียง คุณภาพเสียง และผลของเสียงต่อการได้ยินของมนุษย์ได้"},
    {"code": "7", "text": "อธิบายและประยุกต์ความรู้เรื่องเสียงไปใช้ในชีวิตประจำวันและในเทคโนโลยีต่าง ๆ ได้"},
    {"code": "8", "text": "ทดลองและอธิบายปรากฏการณ์ของแสง ได้แก่ การแทรกสอด การเลี้ยวเบนผ่านสลิตเดี่ยว สลิตคู่ และเกรตติง รวมทั้งการประยุกต์ใช้กับทัศนอุปกรณ์ได้"},
]

# K/P/A summary per ผลการเรียนรู้
KPA = {
    "1": {
        "K": "อธิบายการเกิดคลื่น · ชนิดคลื่น · ปริมาณคลื่น (λ, f, T, v) · ความสัมพันธ์ v = fλ",
        "P": "ทดลองหาความเร็วคลื่น · วิเคราะห์กราฟ x-t, y-t · สื่อสารผลการทดลอง",
        "A": "ทำงานเป็นทีมในห้องปฏิบัติการ · ซื่อสัตย์ทางวิทยาศาสตร์ · ใฝ่เรียนรู้",
    },
    "2": {
        "K": "อธิบายการสะท้อน การหักเห การแทรกสอด การเลี้ยวเบน และคลื่นนิ่ง · กฎของฮอยเกนส์ · v = √(T/μ)",
        "P": "ทดลอง Lab Melde · วาดแผนภาพคลื่น · คำนวณมุมสะท้อน-หักเห · plot กราฟ T² vs L",
        "A": "ความมุ่งมั่นในการทดลอง · ยอมรับผลการวัด · ทำงานกลุ่ม",
    },
    "3": {
        "K": "นิยาม SHM · สมการ T = 2π√(m/k), T = 2π√(L/g) · พลังงานใน SHM · damped + resonance",
        "P": "ทดลองวัด g จากลูกตุ้ม · ทำกราฟ T² vs L · แก้โจทย์ SHM 3 ระดับ",
        "A": "ความรับผิดชอบในทีม · วินัยในการทดลอง · วิเคราะห์ % error อย่างซื่อสัตย์",
    },
    "4": {
        "K": "การเกิดเสียง · การเคลื่อนที่ของเสียง · v = 331 + 0.6T · การกระจัดและคลื่นความดัน",
        "P": "วัดอัตราเร็วเสียงในอากาศ · วิเคราะห์กราฟคลื่นเสียง · เปรียบเทียบเสียงในตัวกลางต่าง ๆ",
        "A": "ใฝ่เรียนรู้ · ทำงานเป็นทีม · ตระหนักผลของเสียง",
    },
    "5": {
        "K": "การสั่นพ้องในท่อ · บีต fb = |f₁ − f₂| · ดอปเพลอร์ f' = f(v ± vₒ)/(v ∓ vₛ)",
        "P": "ทดลอง resonance tube · คำนวณ Doppler 3 ระดับ · STEM ออกแบบเครื่องดนตรี",
        "A": "ความคิดสร้างสรรค์ · ทำงานทีม · ภูมิใจในผลงาน",
    },
    "6": {
        "K": "ระดับความเข้มเสียง β = 10 log(I/I₀) · pitch vs loudness vs timbre · มลพิษทางเสียง",
        "P": "วัดระดับเสียงด้วย sound meter · จำแนกฮาร์มอนิก · ออกแบบมาตรการลดเสียง",
        "A": "ตระหนักผลกระทบของเสียงต่อสุขภาพ · จิตสาธารณะ",
    },
    "7": {
        "K": "การประยุกต์ — โซนาร์ · ultrasound การแพทย์ · เครื่องดนตรี · noise canceling · radar",
        "P": "เชื่อมโยงทฤษฎีกับเทคโนโลยี · นำเสนอ Capstone Project · บูรณาการความรู้",
        "A": "เห็นคุณค่าวิทยาศาสตร์ในชีวิตจริง · ทำงานเป็นทีมระยะยาว",
    },
    "8": {
        "K": "แสงเชิงคลื่น · Young's double slit · single slit · grating · เลนส์ ตา กล้อง · Snell + TIR",
        "P": "ทดลอง Snell + critical angle · คำนวณ d sin θ = mλ · ออกแบบทัศนอุปกรณ์",
        "A": "ใฝ่รู้ในธรรมชาติของแสง · ทำงานเป็นทีม · ปลอดภัยใน Lab",
    },
}

# Units (4 หน่วยการเรียนรู้ตามหลักสูตร · 40 แผน · 80 คาบ)
# Plans 39 (Diagnostic) + 40 (Capstone) จัดอยู่ในหน่วย 4 แสงเชิงคลื่น (ปลายภาคบูรณาการ)
UNITS = [
    {
        "no": 1, "name": "คลื่นกล", "outcomes": ["1", "2"], "plans": 12, "periods": 24, "score": 25,
        "topics": "การเกิดและสมบัติของคลื่น · อัตราเร็วคลื่น · คลื่นต่อเนื่อง+เฟส · การซ้อนทับ+ฮอยเกนส์ · การสะท้อน · การหักเห · การแทรกสอด · การเลี้ยวเบน · คลื่นนิ่งและบูรณาการ · ทบทวน · Lab Melde · Workshop Doppler",
        "plan_range": "1–12",
    },
    {
        "no": 2, "name": "การเคลื่อนที่แบบฮาร์มอนิกอย่างง่าย (SHM)", "outcomes": ["3"], "plans": 8, "periods": 16, "score": 20,
        "topics": "นิยาม SHM · ลูกตุ้มอย่างง่าย · มวลสปริง · เฟส x,v,a · พลังงาน · damped + resonance · Lab วัด g · Workshop ประยุกต์",
        "plan_range": "13–20",
    },
    {
        "no": 3, "name": "เสียงและการได้ยิน", "outcomes": ["4", "5", "6", "7"], "plans": 13, "periods": 26, "score": 30,
        "topics": "การเกิดและการเคลื่อนที่ของเสียง · การกระจัด+คลื่นความดัน · อัตราเร็วเสียง · พฤติกรรมเสียง · ความเข้ม+ระดับเสียง · การได้ยิน+มลพิษ · pitch+คุณภาพเสียง · คลื่นนิ่งในท่อ · บีต · ดอปเพลอร์+คลื่นกระแทก · การประยุกต์ · Lab resonance · STEM เครื่องดนตรี",
        "plan_range": "21–33",
    },
    {
        "no": 4, "name": "แสงเชิงคลื่นและบูรณาการรวม", "outcomes": ["1","2","3","4","5","6","7","8"], "plans": 7, "periods": 14, "score": 25,
        "topics": "แนวคิดแสงเชิงคลื่น+สลิตคู่ · การเลี้ยวเบนผ่านสลิตเดี่ยว · การเลี้ยวเบนผ่านเกรตติง · Lab Snell+TIR · Workshop ทัศนอุปกรณ์ · Two-tier Diagnostic ทบทวนรวม · Capstone นำเสนอโครงงานบูรณาการ",
        "plan_range": "34–40",
    },
]

# Plan list (40 plans)
PLANS = [
    # waves (1-12)
    (1, 1, "คลื่นกล", "การเกิดคลื่นและชนิดของคลื่น", 2, "1", "1,2", "3,4,6", "5E+POE · การสาธิต/ทดลอง · 1, 3"),
    (2, 1, "คลื่นกล", "ความเร็วคลื่น", 2, "1", "1,2,3", "3,4,6", "5E+POE · การทดลอง · 1, 3"),
    (3, 1, "คลื่นกล", "คลื่นดล คลื่นต่อเนื่อง และเฟส", 2, "1", "1,2,5", "3,4,6", "5E+POE · การอภิปรายกลุ่มย่อย · 1"),
    (4, 1, "คลื่นกล", "การซ้อนทับและฮอยเกนส์", 2, "2", "1,2,3", "3,4,6", "5E+POE · การสาธิต+อภิปราย · 1"),
    (5, 1, "คลื่นกล", "การสะท้อนของคลื่น", 2, "2", "1,2,5", "3,4,6", "5E+POE · การทดลอง · 1, 2"),
    (6, 1, "คลื่นกล", "การหักเหของคลื่น", 2, "2", "1,2,3,5", "3,4,6", "5E+POE · การทดลอง · 1, 2"),
    (7, 1, "คลื่นกล", "การแทรกสอดของคลื่น", 2, "2", "1,2,5", "3,4,6", "5E+POE · การทดลอง · 1, 2"),
    (8, 1, "คลื่นกล", "การเลี้ยวเบนของคลื่น", 2, "2", "1,2,5", "3,4,6", "5E+POE · การทดลอง+อภิปราย · 1"),
    (9, 1, "คลื่นกล", "คลื่นนิ่งและบูรณาการ", 2, "2", "1,2,5", "3,4,6", "5E+POE · การทดลอง+โครงงาน · 1"),
    (10, 1, "คลื่นกล", "ทบทวน", 2, "1,2", "1,2,3", "3,4,6", "5E+POE · การถามตอบ+การแก้ปัญหา · 1"),
    (11, 1, "คลื่นกล", "ปฏิบัติการ Melde — คลื่นนิ่งในเส้นเชือก", 2, "2", "1,2,3,4,5", "3,4,6", "5E+POE · การทดลอง (hands-on) · 1, 2"),
    (12, 1, "คลื่นกล", "Workshop แก้โจทย์คลื่นและ Doppler", 2, "2,5", "1,2,3,5", "3,4,6", "5E+POE · Jigsaw+การแก้ปัญหา · 1, 3"),
    # SHM (13-20)
    (13, 2, "SHM", "นิยามและสมการ SHM", 2, "3", "1,2,5", "3,4,6", "5E+POE · การสาธิต+อธิบาย · 1"),
    (14, 2, "SHM", "ลูกตุ้มอย่างง่าย", 2, "3", "1,2,3,5", "3,4,6", "5E+POE · การทดลอง · 1"),
    (15, 2, "SHM", "มวลสปริง", 2, "3", "1,2,3,5", "3,4,6", "5E+POE · การทดลอง · 1"),
    (16, 2, "SHM", "เฟส x, v, a", 2, "3", "1,2,5", "3,4,6", "5E+POE · สถานการณ์จำลอง · 1"),
    (17, 2, "SHM", "พลังงานใน SHM", 2, "3", "1,2,5", "3,4,6", "5E+POE · การอภิปราย+คำนวณ · 1"),
    (18, 2, "SHM", "damped + resonance", 2, "3", "1,2,5", "3,4,6", "5E+POE · กรณีตัวอย่าง · 1, 2"),
    (19, 2, "SHM", "ปฏิบัติการวัดค่า g จากลูกตุ้ม", 2, "3", "1,2,3,4,5", "3,4,6", "5E+POE · การทดลอง (hands-on) · 1, 2"),
    (20, 2, "SHM", "Workshop แก้โจทย์ประยุกต์ SHM", 2, "3", "1,2,3,5", "3,4,6", "5E+POE · Jigsaw+การแก้ปัญหา · 1, 3"),
    # sound (21-33)
    (21, 3, "เสียง", "การเกิดและการเคลื่อนที่ของเสียง", 2, "4", "1,2,5", "3,4,6", "5E+POE · การสาธิต · 1"),
    (22, 3, "เสียง", "การกระจัดและคลื่นความดัน", 2, "4", "1,2,5", "3,4,6", "5E+POE · การอภิปราย · 1"),
    (23, 3, "เสียง", "อัตราเร็วของเสียงและอุณหภูมิ", 2, "4", "1,2,3,5", "3,4,6", "5E+POE · การทดลอง · 1"),
    (24, 3, "เสียง", "พฤติกรรมของคลื่นเสียง", 2, "4,5", "1,2,5", "3,4,6", "5E+POE · การสาธิต+อภิปราย · 1"),
    (25, 3, "เสียง", "ความเข้มและระดับเสียง", 2, "6", "1,2,3", "3,4,6,8", "5E+POE · การทดลอง+คำนวณ · 1, 2"),
    (26, 3, "เสียง", "การได้ยินและมลพิษทางเสียง", 2, "6,7", "1,2,4", "3,4,6,8", "5E+POE · กรณีตัวอย่าง · 1, 2"),
    (27, 3, "เสียง", "ระดับสูงต่ำและคุณภาพเสียง", 2, "5,6", "1,2,5", "3,4,6", "5E+POE · การสาธิต · 1"),
    (28, 3, "เสียง", "คลื่นนิ่งและการสั่นพ้องในท่อ", 2, "5", "1,2,3,5", "3,4,6", "5E+POE · การทดลอง · 1"),
    (29, 3, "เสียง", "บีต", 2, "5", "1,2,5", "3,4,6", "5E+POE · การทดลอง+อภิปราย · 1"),
    (30, 3, "เสียง", "ดอปเพลอร์และคลื่นกระแทก", 2, "5", "1,2,3,5", "3,4,6", "5E+POE · กรณีตัวอย่าง · 1"),
    (31, 3, "เสียง", "การประยุกต์ใช้ความรู้เรื่องเสียง", 2, "7", "1,2,5", "3,4,6,8", "5E+POE · โครงงาน · 1, 4"),
    (32, 3, "เสียง", "ปฏิบัติการ Resonance Tube · วัดความเร็วเสียง", 2, "5", "1,2,3,4,5", "3,4,6", "5E+POE · การทดลอง (hands-on) · 1, 2"),
    (33, 3, "เสียง", "STEM ออกแบบเครื่องดนตรีจากวัสดุเหลือใช้", 2, "5,7", "1,2,3,4,5", "3,4,5,6,8", "5E+STEM · โครงงาน · 1, 2, 4"),
    # light (34-38)
    (34, 4, "แสงเชิงคลื่น", "แนวคิดแสงเชิงคลื่นและสลิตคู่", 2, "8", "1,2,5", "3,4,6", "5E+POE · การสาธิต+ทดลอง · 1"),
    (35, 4, "แสงเชิงคลื่น", "การเลี้ยวเบนของแสงผ่านสลิตเดี่ยว", 2, "8", "1,2,3,5", "3,4,6", "5E+POE · การทดลอง · 1"),
    (36, 4, "แสงเชิงคลื่น", "การเลี้ยวเบนของแสงผ่านเกรตติง", 2, "8", "1,2,3,5", "3,4,6", "5E+POE · การทดลอง · 1"),
    (37, 4, "แสงเชิงคลื่น", "ปฏิบัติการ Snell's Law และ TIR", 2, "8", "1,2,3,4,5", "3,4,6", "5E+POE · การทดลอง (hands-on) · 1, 2"),
    (38, 4, "แสงเชิงคลื่น", "Workshop ทัศนอุปกรณ์ · กล้อง ตา กล้องส่อง", 2, "8", "1,2,3,5", "3,4,6", "5E+POE · Jigsaw+การแก้ปัญหา · 1, 3"),
    # review (39) — อยู่ในหน่วย 4 แสงเชิงคลื่น (ปลายภาคบูรณาการ)
    (39, 4, "แสงเชิงคลื่นและบูรณาการรวม", "Two-tier Diagnostic ทบทวนรวม", 2, "1,2,3,4,5,6,7,8", "1,2,5", "3,4,6", "5E · การวัดมโนทัศน์+การถามตอบ · 1"),
    # capstone (40) — อยู่ในหน่วย 4 แสงเชิงคลื่น
    (40, 4, "แสงเชิงคลื่นและบูรณาการรวม", "นำเสนอโครงงานบูรณาการ (Capstone)", 2, "1,2,3,4,5,6,7,8", "1,2,3,4,5", "3,4,6,8", "PBL+5E · โครงงาน+การนำเสนอ · 1, 2, 3, 4"),
]

# ─── ไฟล์ 04 — ตารางการวิเคราะห์มาตรฐาน ตัวชี้วัด ────────────
def build_04():
    doc = Document()
    set_landscape(doc)
    add_title(doc, "ตาราง การวิเคราะห์มาตรฐาน ตัวชี้วัดและผลการเรียนรู้")
    add_header_block(doc)
    add_para(doc, "", size=8, space_after=4)

    headers = [
        "ผลการเรียนรู้",
        "เนื้อหา/สาระการเรียนรู้",
        "K (ความรู้)",
        "P (ทักษะกระบวนการ)",
        "A (เจตคติ)",
        "คุณลักษณะ\nอันพึงประสงค์",
        "สมรรถนะสำคัญ\nของผู้เรียน",
        "บูรณาการหลักสูตร\n/กิจกรรมสถานศึกษา",
    ]
    widths = [3.2, 4.0, 3.8, 3.8, 3.0, 2.6, 2.6, 2.6]
    t = doc.add_table(rows=1+len(LEARN_OUTCOMES), cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    for i, w in enumerate(widths):
        for cell in t.columns[i].cells:
            cell.width = Cm(w)
    # header
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, size=12, bold=True,
                 color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER,
                 shade_color="1f4e79")
    # rows
    for ri, lo in enumerate(LEARN_OUTCOMES, 1):
        # find topics for this outcome
        topics = []
        for u in UNITS:
            if lo["code"] in u["outcomes"]:
                topics.append(f"({u['no']}) {u['name']}")
        topics_text = " · ".join(topics) if topics else "—"

        kpa = KPA.get(lo["code"], {"K":"", "P":"", "A":""})
        # คุณลักษณะ + สมรรถนะ + บูรณาการ (มาตรฐาน)
        attr = "มีวินัย\nใฝ่เรียนรู้\nมุ่งมั่นในการทำงาน"
        comp = "การสื่อสาร\nการคิด\nการแก้ปัญหา\nการใช้เทคโนโลยี"
        integ = "1. สตรีวิทยาศึกษา\n3. เศรษฐกิจพอเพียง"
        # ปรับสำหรับบางผลฯ
        if lo["code"] in ("6","7"):
            attr += "\nจิตสาธารณะ"
            integ = "1. สตรีวิทยาศึกษา\n2. สตรีวิทยารวมพลังรักษ์ พิทักษ์โลก\n3. เศรษฐกิจพอเพียง"
        if lo["code"] == "7":
            integ += "\n4. บูรณาการอาชีพ"
        cells = t.rows[ri].cells
        set_cell(cells[0], f"ผลการเรียนรู้ที่ {lo['code']}\n{lo['text']}",
                 size=11.5, vertical=WD_ALIGN_VERTICAL.TOP)
        set_cell(cells[1], topics_text, size=11.5, vertical=WD_ALIGN_VERTICAL.TOP)
        set_cell(cells[2], kpa["K"], size=11.5, vertical=WD_ALIGN_VERTICAL.TOP)
        set_cell(cells[3], kpa["P"], size=11.5, vertical=WD_ALIGN_VERTICAL.TOP)
        set_cell(cells[4], kpa["A"], size=11.5, vertical=WD_ALIGN_VERTICAL.TOP)
        set_cell(cells[5], attr, size=11.5, vertical=WD_ALIGN_VERTICAL.TOP)
        set_cell(cells[6], comp, size=11.5, vertical=WD_ALIGN_VERTICAL.TOP)
        set_cell(cells[7], integ, size=11.5, vertical=WD_ALIGN_VERTICAL.TOP)

    add_signatures(doc)
    add_para(doc, "หมายเหตุ บูรณาการหลักสูตรสถานศึกษา/กิจกรรมสถานศึกษา:  1. สตรีวิทยาศึกษา   2. สตรีวิทยารวมพลังรักษ์ พิทักษ์โลก   3. เศรษฐกิจพอเพียง   4. บูรณาการอาชีพ",
             size=11, italic=True, color=(0x55,0x55,0x55), space_before=12)

    out = OUTDIR / "04_ตารางการวิเคราะห์มาตรฐาน_ตัวชี้วัด.docx"
    doc.save(str(out)); print(f"  ✅ {out.name}")
    return out

# ─── ไฟล์ 05 — ตารางการออกแบบหน่วยการเรียนรู้ ────────────────
def build_05():
    doc = Document()
    set_landscape(doc)
    add_title(doc, "การออกแบบหน่วยการเรียนรู้")
    add_header_block(doc)
    add_para(doc, "", size=8, space_after=4)

    headers = [
        "ชื่อหน่วย\nการเรียนรู้",
        "ผลการเรียนรู้",
        "สาระการเรียนรู้\n/เนื้อหา",
        "รูปแบบ/เทคนิค\nการจัดกิจกรรม\nและสื่อ",
        "ชิ้นงาน/ภาระงาน",
        "เครื่องมือวัด",
        "จำนวน\nคาบ",
        "คะแนน",
        "แผนการ\nจัดการ\nเรียนรู้ที่",
    ]
    widths = [2.5, 2.5, 4.0, 4.5, 3.5, 3.0, 1.2, 1.2, 1.6]
    t = doc.add_table(rows=1+len(UNITS), cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    for i, w in enumerate(widths):
        for cell in t.columns[i].cells:
            cell.width = Cm(w)
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, size=12, bold=True,
                 color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER,
                 shade_color="1f4e79")

    for ri, u in enumerate(UNITS, 1):
        cells = t.rows[ri].cells
        outc = "ผลการเรียนรู้ที่\n" + ", ".join(u["outcomes"])
        # รูปแบบ/เทคนิค + สื่อ
        form_lines = ["รูปแบบ/เทคนิค:",
                      "• สืบเสาะหาความรู้ 5E + POE",
                      "• การทดลอง/ปฏิบัติการ Lab",
                      "• Concept Cartoon + CER"]
        if u["no"] == 3:
            form_lines.append("• STEM Engineering Design")
        if u["no"] == 4:
            form_lines.append("• Two-tier Diagnostic Test")
            form_lines.append("• Project-Based Learning (Capstone)")
            form_lines.append("• Gallery Walk + Peer Review")
        form_lines += ["",
                       "สื่อการจัดการเรียนรู้:",
                       "• Infographic แผนการสอน",
                       "• POE Sheet / ใบงาน Calc / Spot the Error",
                       "• Four-tier Diagnostic (Apps Script)",
                       "• PhET / Simulation interactive",
                       "• อุปกรณ์ทดลองจริง (Lab kit)"]
        form_text = "\n".join(form_lines)

        # ชิ้นงาน/ภาระงาน
        work_lines = ["• ใบบันทึก POE 3 ฐาน",
                      "• ใบงานคำนวณ Calc",
                      "• ใบงาน Spot the Error",
                      "• Metacognitive Journal 3-2-1"]
        if u["no"] == 3:
            work_lines.append("• เครื่องดนตรี STEM (ชิ้นงาน)")
        if u["no"] == 4:
            work_lines += ["• Two-tier Test 20 ข้อ + Personal Learning Plan",
                           "• ชิ้นงาน/Demo Capstone + Presentation 10 นาที"]
        work_text = "\n".join(work_lines)

        # เครื่องมือวัด
        tool_lines = ["• Four-tier Pre/Post (F-test)",
                      "• POE Rubric (0–3)",
                      "• แบบสังเกตพฤติกรรม (OB)",
                      "• Traffic Light Card",
                      "• MJ 3-2-1 Reflection"]
        if u["no"] == 4:
            tool_lines += ["• Capstone Rubric 5 มิติ + Peer Evaluation",
                           "• Hake gain + Diagnostic Report"]
        tool_text = "\n".join(tool_lines)

        set_cell(cells[0], f"หน่วยที่ {u['no']}\n{u['name']}",
                 size=11.5, vertical=WD_ALIGN_VERTICAL.TOP, bold=True)
        set_cell(cells[1], outc, size=11.5, vertical=WD_ALIGN_VERTICAL.TOP)
        set_cell(cells[2], u["topics"], size=11, vertical=WD_ALIGN_VERTICAL.TOP)
        set_cell(cells[3], form_text, size=11, vertical=WD_ALIGN_VERTICAL.TOP)
        set_cell(cells[4], work_text, size=11, vertical=WD_ALIGN_VERTICAL.TOP)
        set_cell(cells[5], tool_text, size=11, vertical=WD_ALIGN_VERTICAL.TOP)
        set_cell(cells[6], str(u["periods"]), size=12,
                 align=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_VERTICAL.CENTER, bold=True)
        set_cell(cells[7], str(u["score"]), size=12,
                 align=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_VERTICAL.CENTER, bold=True)
        set_cell(cells[8], u["plan_range"], size=12,
                 align=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_VERTICAL.CENTER, bold=True)

    # รวม
    total_periods = sum(u["periods"] for u in UNITS)
    total_score = sum(u["score"] for u in UNITS)
    add_para(doc, f"รวม  {sum(u['plans'] for u in UNITS)} แผน  ·  {total_periods} คาบ  ·  {total_score} คะแนน",
             size=14, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=8)

    add_signatures(doc)

    # หมายเหตุ — รูปแบบการสอน
    add_para(doc, "", size=8, space_before=8)
    add_para(doc, "หมายเหตุ  รูปแบบ/วิธีการจัดกิจกรรมการเรียนการสอนที่ครูใช้",
             size=13, bold=True, space_before=6)
    # 2-column checklist
    methods = [
        ("☑ การอธิบาย", "☑ การสืบสวนสอบสวน"),
        ("☑ การสาธิต/ทดลอง", "☑ กลุ่มสืบค้นความรู้"),
        ("☐ การใช้เกมประกอบ", "☑ กลุ่มสัมพันธ์"),
        ("☑ สถานการณ์จำลอง", "☑ การเรียนรู้แบบร่วมมือ"),
        ("☑ กรณีตัวอย่าง", "☑ ความคิดรวบยอด"),
        ("☐ บทบาทสมมุติ", "☐ อริยสัจ 4"),
        ("☑ การแก้ไขสถานการณ์", "☑ การศึกษาค้นคว้าด้วยตนเอง"),
        ("☐ โปรแกรมสำเร็จรูป", "☐ การทัศนะศึกษานอกสถานที่"),
        ("☐ ศูนย์การเรียน", "☑ การเรียนรู้จากห้องสมุด"),
        ("☐ ชุดการสอน", "☑ การพัฒนากระบวนการคิด"),
        ("☑ คอมพิวเตอร์ช่วยสอน", "☐ การใช้ภูมิปัญญาท้องถิ่น"),
        ("☑ โครงงาน", "☑ การอภิปรายกลุ่มย่อย"),
        ("☑ การถามตอบ", "☑ การแก้ปัญหา"),
        ("☑ การอธิบายประกอบสื่อนำเสนอ", "☑ อื่น ๆ ระบุ 5E+POE / STEM / PBL"),
    ]
    mt = doc.add_table(rows=len(methods), cols=2)
    mt.autofit = False
    for col in mt.columns:
        col.width = Cm(8)
        for cell in col.cells:
            cell.width = Cm(8)
    for ri, (left, right) in enumerate(methods):
        set_cell(mt.rows[ri].cells[0], left, size=12, vertical=WD_ALIGN_VERTICAL.CENTER)
        set_cell(mt.rows[ri].cells[1], right, size=12, vertical=WD_ALIGN_VERTICAL.CENTER)
    # remove borders
    for cell in mt._tbl.iter(qn("w:tc")):
        tcPr = cell.find(qn("w:tcPr"))
        if tcPr is not None:
            for tag in list(tcPr):
                if tag.tag == qn("w:tcBorders"):
                    tcPr.remove(tag)
            tcBorders = OxmlElement("w:tcBorders")
            for edge in ("top","left","bottom","right"):
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"), "nil")
                tcBorders.append(b)
            tcPr.append(tcBorders)

    out = OUTDIR / "05_ตารางการออกแบบหน่วยการเรียนรู้.docx"
    doc.save(str(out)); print(f"  ✅ {out.name}")
    return out

# ─── ไฟล์ 08 — ตารางแผนการจัดการเรียนรู้ ─────────────────────
def build_08():
    doc = Document()
    set_landscape(doc)
    add_title(doc, "ตารางแผนการจัดการเรียนรู้")
    # custom header for 08 (slightly different)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(8)
    line = "รายวิชา ฟิสิกส์ 3 (ว 30203)   ชั้นมัธยมศึกษาปีที่ 5   จำนวน 2 หน่วยกิต   ภาคเรียนที่ 1   ปีการศึกษา 2569"
    r = p.add_run(line); set_font(r, size=14)

    headers = [
        "หน่วย\nที่", "ชื่อหน่วย", "เรื่อง", "จำนวน\nคาบ",
        "ผลการ\nเรียนรู้", "สมรรถนะ\nของผู้เรียน", "คุณลักษณะ\nอันพึงประสงค์", "รูปแบบ/เทคนิค\nการจัดกิจกรรม/บูรณาการ",
    ]
    widths = [1.2, 2.8, 7.5, 1.2, 2.0, 2.0, 2.4, 5.0]
    t = doc.add_table(rows=1+len(PLANS), cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    for i, w in enumerate(widths):
        for cell in t.columns[i].cells:
            cell.width = Cm(w)
    for i, h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h, size=12, bold=True,
                 color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER,
                 shade_color="1f4e79")

    for ri, plan in enumerate(PLANS, 1):
        plan_no, unit_no, unit_name, topic, periods, outcomes, comp, attr, form = plan
        cells = t.rows[ri].cells
        set_cell(cells[0], str(unit_no), size=12,
                 align=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_VERTICAL.CENTER, bold=True)
        set_cell(cells[1], unit_name, size=11.5,
                 align=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_VERTICAL.CENTER)
        set_cell(cells[2], f"แผนการจัดการเรียนรู้ที่ {plan_no}\nเรื่อง  {topic}",
                 size=11.5, vertical=WD_ALIGN_VERTICAL.CENTER)
        set_cell(cells[3], str(periods), size=12,
                 align=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_VERTICAL.CENTER, bold=True)
        set_cell(cells[4], outcomes, size=11.5,
                 align=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_VERTICAL.CENTER)
        set_cell(cells[5], comp, size=11.5,
                 align=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_VERTICAL.CENTER)
        set_cell(cells[6], attr, size=11.5,
                 align=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_VERTICAL.CENTER)
        set_cell(cells[7], form, size=11,
                 vertical=WD_ALIGN_VERTICAL.CENTER)

    # รวม
    total_periods = sum(p[4] for p in PLANS)
    add_para(doc, f"รวม  {len(PLANS)} แผน  ·  {total_periods} คาบ",
             size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)

    add_signatures(doc)

    # หมายเหตุ
    add_para(doc, "", size=8, space_before=8)
    add_para(doc, "หมายเหตุ", size=13, bold=True, space_before=6)

    # 2-column tables for สมรรถนะ + คุณลักษณะ
    add_para(doc, "สมรรถนะของผู้เรียน", size=12, bold=True, space_before=4)
    comp_data = [
        ("1", "ความสามารถในการสื่อสาร"),
        ("2", "ความสามารถในการคิด"),
        ("3", "ความสามารถในการแก้ปัญหา"),
        ("4", "ความสามารถในการใช้ทักษะชีวิต"),
        ("5", "ความสามารถในการใช้เทคโนโลยี"),
    ]
    ct = doc.add_table(rows=1+len(comp_data), cols=2)
    ct.style = "Table Grid"
    ct.autofit = False
    ct.columns[0].width = Cm(2.0)
    ct.columns[1].width = Cm(10.0)
    set_cell(ct.rows[0].cells[0], "ข้อที่", size=12, bold=True,
             color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER, shade_color="1f4e79")
    set_cell(ct.rows[0].cells[1], "สมรรถนะของผู้เรียน", size=12, bold=True,
             color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER, shade_color="1f4e79")
    for i, (n, t_) in enumerate(comp_data, 1):
        set_cell(ct.rows[i].cells[0], n, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(ct.rows[i].cells[1], t_, size=12)

    add_para(doc, "", size=8)
    add_para(doc, "คุณลักษณะอันพึงประสงค์", size=12, bold=True, space_before=4)
    attr_data = [
        ("1", "รักชาติ ศาสน์ กษัตริย์"),
        ("2", "ซื่อสัตย์สุจริต"),
        ("3", "มีวินัย"),
        ("4", "ใฝ่เรียนรู้"),
        ("5", "อยู่อย่างพอเพียง"),
        ("6", "มุ่งมั่นในการทำงาน"),
        ("7", "รักความเป็นไทย"),
        ("8", "มีจิตสาธารณะ"),
    ]
    at = doc.add_table(rows=1+len(attr_data), cols=2)
    at.style = "Table Grid"
    at.autofit = False
    at.columns[0].width = Cm(2.0)
    at.columns[1].width = Cm(10.0)
    set_cell(at.rows[0].cells[0], "ข้อที่", size=12, bold=True,
             color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER, shade_color="1f4e79")
    set_cell(at.rows[0].cells[1], "คุณลักษณะอันพึงประสงค์", size=12, bold=True,
             color=(255,255,255), align=WD_ALIGN_PARAGRAPH.CENTER, shade_color="1f4e79")
    for i, (n, t_) in enumerate(attr_data, 1):
        set_cell(at.rows[i].cells[0], n, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell(at.rows[i].cells[1], t_, size=12)

    out = OUTDIR / "08_ตารางแผนการจัดการเรียนรู้.docx"
    doc.save(str(out)); print(f"  ✅ {out.name}")
    return out

def main():
    OUTDIR.mkdir(parents=True, exist_ok=True)
    print("─── สร้างเอกสารประกอบ 3 ไฟล์ ───")
    build_04()
    build_05()
    build_08()
    print(f"\nเก็บที่: {OUTDIR}")

if __name__ == "__main__":
    main()
