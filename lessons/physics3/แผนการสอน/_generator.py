#!/usr/bin/env python3
"""
Generator v2 · ฟิสิกส์ 3 ม.5
แปลง Infographic HTML → DOCX แบบบรรยายความ + วPA เชี่ยวชาญ
ใช้รูปแบบ astronomy COSMOS LOG (Thai numerals · TH SarabunPSK · ตารางสรุป · บรรยายครู/นักเรียน/เป้าหมาย)
"""
import os, re, sys
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pythainlp.tokenize import word_tokenize

ZWSP = "​"
THAI_NUMS = "๐๑๒๓๔๕๖๗๘๙"
def thai_num(n):
    """แปลงเลขอารบิก → เลขไทย"""
    return "".join(THAI_NUMS[int(d)] if d.isdigit() else d for d in str(n))

def strip_html(s):
    """ลบ HTML tags จาก string"""
    if not s: return ""
    return re.sub(r"<[^>]+>", "", s)

BASE = Path("/Users/komanepapato/Documents/วิจัย/wave-mechanics-research/lessons/physics3")
OUTDIR = BASE / "แผนการสอน"

UNITS = [
    ("waves",    "หน่วยที่ 1 คลื่นกล",           "บทที่ 9 คลื่นกล",      "ม.5", "ว 5.2 ม.5/1-5",  12),
    ("SHM",      "หน่วยที่ 2 การเคลื่อนที่แบบฮาร์มอนิกอย่างง่าย", "บทที่ 11 SHM",       "ม.5", "ว 5.2 ม.5/3",     8),
    ("sound",    "หน่วยที่ 3 เสียง",             "บทที่ 12 เสียง",      "ม.5", "ว 5.2 ม.5/4-7",  13),
    ("light",    "หน่วยที่ 4 แสงเชิงคลื่น",       "บทที่ 10 แสงเชิงคลื่น", "ม.5", "ว 2.3 ม.5",      5),
    ("review",   "หน่วยทบทวนรวม",                 "Two-tier Diagnostic", "ม.5", "ทบทวน ว 5.2 + ว 2.3", 1),
    ("capstone", "หน่วย Capstone Project",        "บูรณาการตลอดเทอม",   "ม.5", "บูรณาการ ว 5.2 + ว 2.3", 1),
]

FONT = "TH SarabunPSK"

# ─── มาตรฐานขนาดและการย่อหน้า (Layout System) ────────────────
SZ_H1 = 18      # หัวข้อใหญ่ ๑. ๒. ๓. ...
SZ_H2 = 16      # หัวข้อกิจกรรม ๗.X
SZ_H3 = 14      # หัวข้อย่อย ๖.๑ ๖.๒ + bold sub-heading
SZ_BODY = 14    # ย่อหน้าหลัก
SZ_BULLET = 14  # bullet/numbered
SZ_TABLE = 12   # ข้อความในตาราง
SZ_TAGS = 12    # tag line [วPA: ...]
SZ_NOTE = 11    # หมายเหตุเล็ก

# Indent system (cm) — เป็นมาตรฐานเดียวทั้งเอกสาร
IND_BODY = 0.6        # ย่อหน้าหลัก (เพิ่มจาก 0.3 → 0.6 ให้โผล่ชัด)
IND_BODY_FL = 0.8     # first-line indent ของย่อหน้าบรรยาย
IND_H3 = 0.3          # หัวข้อย่อย
IND_BULLET = 1.2      # bullet
IND_BULLET_FL = -0.5  # hanging indent ของ bullet
IND_NUMBER = 1.2      # numbered
IND_NUMBER_FL = -0.7  # hanging indent ของ numbered

# ─────────────────────────────────────────────────────────────
# Thai line-break helper
def th_break(s):
    if not s: return ""
    try:
        toks = word_tokenize(s, engine="newmm", keep_whitespace=True)
        return ZWSP.join(toks)
    except Exception:
        return s

# ─────────────────────────────────────────────────────────────
# Style helpers
def set_th_font(run, size=14, bold=False, color=None, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rfont = rpr.find(qn("w:rFonts"))
    if rfont is None:
        rfont = OxmlElement("w:rFonts")
        rpr.append(rfont)
    for attr in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"):
        rfont.set(qn(attr), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(doc, text, *, size=14, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, color=None,
             space_before=0, space_after=2, indent_left=0,
             first_line_indent=0, line_spacing=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent_left:
        p.paragraph_format.left_indent = Cm(indent_left)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    if line_spacing:
        p.paragraph_format.line_spacing = line_spacing
    r = p.add_run(th_break(text))
    set_th_font(r, size=size, bold=bold, color=color, italic=italic)
    return p

def add_h1(doc, text, *, color=(0x1f,0x2d,0x5c)):
    """หัวข้อใหญ่ ๑. ๒. ๓. ..."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm = OxmlElement("w:bottom")
    btm.set(qn("w:val"), "single")
    btm.set(qn("w:sz"), "12")
    btm.set(qn("w:space"), "1")
    btm.set(qn("w:color"), "1f2d5c")
    pBdr.append(btm)
    pPr.append(pBdr)
    r = p.add_run(th_break(text))
    set_th_font(r, size=SZ_H1, bold=True, color=color)
    return p

def add_h2(doc, text, *, color=(0x2c,0x5f,0xa0)):
    """หัวข้อกิจกรรม ๗.X — ขั้น 5E"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(th_break(text))
    set_th_font(r, size=SZ_H2, bold=True, color=color)
    return p

def add_h3(doc, text, *, indent=IND_H3):
    """หัวข้อย่อย ๖.๑ ๖.๒ และ subheading bold"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(th_break(text))
    set_th_font(r, size=SZ_H3, bold=True, color=(0x2d,0x3d,0x5a))
    return p

def add_bullet(doc, text, *, size=SZ_BULLET):
    """bullet มาตรฐาน · indent 1.2cm · hanging -0.5"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(IND_BULLET)
    p.paragraph_format.first_line_indent = Cm(IND_BULLET_FL)
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run("•  " + th_break(text))
    set_th_font(r, size=size)
    return p

def add_numbered(doc, num, text, *, size=SZ_BULLET):
    """numbered มาตรฐาน · indent 1.2cm · hanging -0.7"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(IND_NUMBER)
    p.paragraph_format.first_line_indent = Cm(IND_NUMBER_FL)
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run(f"{num}.  " + th_break(text))
    set_th_font(r, size=size)
    return p

def set_cell_shading(cell, hex_color):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(sh)

def set_cell(cell, text, *, size=12, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             color=None, shade=None, vertical=WD_ALIGN_VERTICAL.TOP):
    cell.text = ""
    cell.vertical_alignment = vertical
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(th_break(str(text)))
    set_th_font(r, size=size, bold=bold, color=color)
    if shade:
        set_cell_shading(cell, shade)

def add_table(doc, headers, rows, col_widths_cm=None, header_shade="1f4e79"):
    """ตารางมาตรฐาน · header สีน้ำเงิน + ลายเส้นเทา"""
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = False
    if col_widths_cm:
        for i,w in enumerate(col_widths_cm):
            for cell in t.columns[i].cells:
                cell.width = Cm(w)
    # header
    for i,h in enumerate(headers):
        set_cell(t.rows[0].cells[i], h,
                 size=12.5, bold=True, color=(255,255,255),
                 align=WD_ALIGN_PARAGRAPH.CENTER,
                 shade=header_shade,
                 vertical=WD_ALIGN_VERTICAL.CENTER)
    # body
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            set_cell(t.rows[ri+1].cells[ci], val, size=12,
                     vertical=WD_ALIGN_VERTICAL.TOP)
    return t

def add_divider(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("◆ ◆ ◆")
    set_th_font(r, size=11, color=(0x99,0x99,0x99))

def add_pagebreak(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)

# ─────────────────────────────────────────────────────────────
# Infographic HTML parser (เหมือนเดิม)
def parse_infographic(html_path):
    soup = BeautifulSoup(Path(html_path).read_text(encoding="utf-8"), "html.parser")
    data = {}
    hero = soup.find(class_="hero")
    if hero:
        num = hero.find(class_="plan-num")
        data["plan_num"] = num.get_text(strip=True) if num else ""
        h1 = hero.find("h1")
        data["title"] = h1.get_text(strip=True) if h1 else ""
        sub = hero.find(class_="subtitle")
        data["subtitle"] = sub.get_text(strip=True) if sub else ""
        course_p = hero.find("p")
        data["course_info"] = course_p.get_text(strip=True) if course_p else ""
        data["tags"] = [t.get_text(strip=True) for t in hero.find_all(class_="tag")]
        data["meta"] = [m.get_text(" ", strip=True) for m in hero.find_all(class_="meta-item")]
    cb = soup.find(class_="concept-box")
    if cb:
        for br in cb.find_all("br"): br.replace_with("\n")
        data["concept"] = cb.get_text(" ", strip=False).strip()
    data["refr_cards"] = []
    for c in soup.find_all(class_="refr-card"):
        title = c.find(class_="refr-card-title")
        sub = c.find(class_="refr-card-sub")
        note = c.find(class_="refr-card-note")
        data["refr_cards"].append({
            "title": title.get_text(strip=True) if title else "",
            "sub": sub.get_text(strip=True) if sub else "",
            "note": note.get_text(strip=True) if note else "",
        })
    fb = soup.find(class_="formula-box")
    if fb:
        rows = []
        for row in fb.find_all(class_="formula-row"):
            tag = row.find(class_="formula-tag")
            formula = row.find(class_="formula")
            note = ""
            for s in row.find_all("span"):
                cls = " ".join(s.get("class") or [])
                if not cls:
                    note = s.get_text(strip=True); break
            rows.append({
                "tag": tag.get_text(strip=True) if tag else "",
                "formula": formula.get_text(strip=True) if formula else "",
                "note": note,
            })
        data["formulas"] = rows
    data["misc"] = []
    mb = soup.find(class_="misc-box")
    if mb:
        for item in mb.find_all(class_="misc-item"):
            mid = item.find(class_="misc-id")
            text = item.find(class_="misc-text")
            correct = item.find(class_="misc-correct")
            data["misc"].append({
                "id": mid.get_text(strip=True) if mid else "",
                "wrong": text.get_text(" ", strip=True) if text else "",
                "correct": correct.get_text(" ", strip=True) if correct else "",
            })
    data["K"] = []; data["P"] = []; data["A"] = []
    for grp in soup.find_all(class_="obj-group"):
        otype = grp.find(class_="obj-type")
        if not otype: continue
        clsstr = " ".join(otype.get("class") or [])
        for li in grp.find_all("li"):
            code = li.find(class_="obj-code")
            code_text = code.get_text(strip=True) if code else ""
            full = li.get_text(" ", strip=True)
            if code_text: full = full.replace(code_text, "", 1).strip()
            entry = {"code": code_text, "text": full}
            if " k" in clsstr or clsstr.endswith("k"): data["K"].append(entry)
            elif " p" in clsstr or clsstr.endswith("p"): data["P"].append(entry)
            elif " a" in clsstr or clsstr.endswith("a"): data["A"].append(entry)
    data["phases"] = []
    for ph in soup.find_all(class_="phase"):
        dot = ph.find(class_="phase-dot")
        name = ph.find(class_="phase-name")
        time = ph.find(class_="phase-time")
        objb = ph.find_all(class_="obj-badge")
        miscb = ph.find_all(class_="misc-badge")
        vpab = ph.find_all(class_="vpa-badge")
        poe = ph.find(class_="poe-chip")
        actbadge = ph.find(class_="activity-badge")
        lis = []
        for li in ph.find_all("li"):
            t = li.find(class_="li-text")
            if t: lis.append(t.get_text(" ", strip=True))
        data["phases"].append({
            "code": dot.get_text(strip=True) if dot else "",
            "name": name.get_text(strip=True) if name else "",
            "time": time.get_text(strip=True) if time else "",
            "obj": [b.get_text(strip=True) for b in objb],
            "misc": [b.get_text(strip=True) for b in miscb],
            "vpa": [b.get_text(strip=True) for b in vpab],
            "poe": poe.get_text(strip=True) if poe else "",
            "activity": actbadge.get_text(strip=True) if actbadge else "",
            "activities": lis,
        })
    data["tools"] = []
    for ti in soup.find_all(class_="tool-item"):
        num = ti.find(class_="tool-num")
        full = ti.get_text(" ", strip=True)
        if num:
            full = full.replace(num.get_text(strip=True), "", 1).strip()
        data["tools"].append(full)
    data["assess"] = []
    at = soup.find(class_="assess-table")
    if at:
        for tr in at.find_all("tr")[1:]:
            tds = [td.get_text(" ", strip=True) for td in tr.find_all("td")]
            if tds: data["assess"].append(tds)
    data["vpa"] = []
    vt = soup.find(class_="vpa-table")
    if vt:
        for tr in vt.find_all("tr")[1:]:
            tds = [td.get_text(" ", strip=True) for td in tr.find_all("td")]
            if tds and len(tds) >= 3: data["vpa"].append(tds)
    data["rubric"] = []
    for rs in soup.find_all(class_="rubric-score"):
        num = rs.find(class_="score-num")
        desc = rs.find(class_="score-desc")
        data["rubric"].append({
            "score": num.get_text(strip=True) if num else "",
            "desc": desc.get_text(strip=True) if desc else "",
        })
    sim = soup.find(class_="sim-banner")
    if sim:
        title = sim.find(class_="sim-banner-title")
        sub = sim.find(class_="sim-banner-sub")
        data["sim"] = {
            "title": title.get_text(strip=True) if title else "",
            "url": sub.get_text(strip=True) if sub else "",
            "badges": [b.get_text(strip=True) for b in sim.find_all(class_="sim-badge")],
        }
    return data

# ─────────────────────────────────────────────────────────────
# Phase narrative templates (ครู / นักเรียน / เป้าหมาย)
PHASE_INFO = {
    "TLC": {
        "icon": "🚦",
        "name": "Traffic Light Card · ตรวจสอบสภาพการเรียนรู้ก่อนเรียน",
        "teacher": "ครูเปิดประเด็นด้วยคำถามชวนคิดที่ใกล้ตัวนักเรียน ใช้ Traffic Light Card เก็บข้อมูลความเข้าใจภาพรวมของห้อง จากนั้นประกาศเป้าหมายของคาบและเกณฑ์การประเมินให้นักเรียนเห็นชัดเจน (Visible Learning)",
        "student": "นักเรียนรายบุคคลชู 🟢 (เข้าใจ) · 🟡 (ไม่แน่ใจ) · 🔴 (ไม่เข้าใจ) ตอบคำถามตรงหน้า แล้วบันทึกความรู้สึกของตนเองลงใน TL-card ช่อง Pre พร้อมระบุเหตุผลสั้น ๆ 1-2 ประโยค",
        "goal": "ใช้เป็น formative pre-assessment ของคาบ — ครูเห็น baseline ของห้องและปรับการสอนได้ทันที · สอดคล้อง วPA ตัวชี้วัด ๔ (กระตุ้นแรงจูงใจ) และ ๗ (Self-Regulated Learning)",
    },
    "ACT": {
        "icon": "⚡",
        "name": "Activate Prior Knowledge · ทบทวนความรู้เดิม",
        "teacher": "ครูใช้ Kahoot Quick Quiz หรือคำถามแบบเร็ว เพื่อทบทวนความรู้จากแผนก่อนหน้า · ครูสังเกตคะแนนกลุ่มและระบุข้อที่ตอบผิดมากที่สุดเพื่อย้ำในขั้น Explain ภายหลัง",
        "student": "นักเรียนรายบุคคลตอบ Kahoot บนอุปกรณ์ของตน · จดข้อที่ยังไม่แน่ใจไว้ในสมุด · แลกเปลี่ยนคำตอบกับเพื่อนคู่ข้างถ้าครูเปิดให้ปรึกษา",
        "goal": "เชื่อมโยงความรู้เดิมกับการเรียนรู้ใหม่ตามแนวคิด constructivist · ครูได้ baseline ของห้องเพื่อปรับแผน · สอดคล้อง วPA ตัวชี้วัด ๒ (เชื่อมโยงความรู้/ประสบการณ์เดิม) โดยตรง",
    },
    "F-Pre": {
        "icon": "📝",
        "name": "Formative Pre-test · Four-tier Diagnostic",
        "teacher": "ครูเปิดข้อสอบ Four-tier Diagnostic บนระบบ Apps Script ที่วัดทั้ง<strong>ความเข้าใจมโนทัศน์</strong>และ<strong>ความมั่นใจ</strong>ของนักเรียน · ครูเฝ้าดู dashboard เพื่อเห็น baseline ของ Sound concept / Misconception / Guessing / No-knowledge",
        "student": "นักเรียนรายบุคคลทำข้อสอบ Pre-test 5 ข้อบนอุปกรณ์ · ทุกข้อต้องเลือก Tier 1 คำตอบ + Tier 2 เหตุผล + Tier 3 ระดับความมั่นใจ + Tier 4 misconception code ระบบส่งข้อมูลขึ้น Apps Script อัตโนมัติ",
        "goal": "เก็บ baseline เชิงมโนทัศน์ก่อนเรียน เพื่อเปรียบเทียบกับ Post-test หลังจบคาบ (วัด Hake gain) · สอดคล้อง วPA ตัวชี้วัด ๒ (เชื่อมโยงความรู้เดิม) และ ๖ (ข้อมูลย้อนกลับ)",
    },
    "E1": {
        "icon": "🎯",
        "name": "Engage · สร้างความสนใจและขัดแย้งทางมโนทัศน์",
        "teacher": "ครูใช้สื่อ Concept Cartoon (หรือ stimulus อื่น เช่น วิดีโอ ปรากฏการณ์จริง) ที่มีตัวการ์ตูนถกเถียงด้วยความเชื่อต่าง ๆ บางความเชื่อตรงกับ misconception เป้าหมาย · ครูทำหน้าที่ผู้ดำเนินรายการ ไม่เฉลย",
        "student": "นักเรียนรายบุคคลเลือกข้างและให้เหตุผลแบบ CER (Claim-Evidence-Reasoning) เบื้องต้น · เขียน Predict (P ใน POE) ลงในใบบันทึก · จับคู่ Think-Pair-Share แลกเหตุผลกับเพื่อน · ครูสุ่มถามแบบ no-hands-up",
        "goal": "กระตุ้นการเรียนรู้ด้วยการขัดแย้งทางมโนทัศน์ (Cognitive Conflict) นักเรียนเปิดเผยความเข้าใจเดิม-ทั้งที่ถูกและคลาดเคลื่อน · สอดคล้อง วPA ตัวชี้วัด ๑ (เข้าถึงสิ่งที่เรียน) และ ๔ (กระตุ้นแรงจูงใจ)",
    },
    "E2": {
        "icon": "🔬",
        "name": "Explore · ปฏิบัติการ/สำรวจด้วยตนเอง",
        "teacher": "ครูจัดฐานปฏิบัติการหรือกิจกรรมสำรวจ · ทำหน้าที่เป็นโค้ช เดินดูทุกกลุ่ม ถามคำถามนำให้นักเรียนคิด ไม่เฉลยทันที · บันทึกพฤติกรรมการทำงานของนักเรียนใน OB sheet",
        "student": "นักเรียนกลุ่ม 4 คนทำกิจกรรมตามฐานต่าง ๆ · แบ่งบทบาทกันชัดเจน (ผู้วัด/ผู้บันทึก/ผู้คำนวณ/ผู้นำเสนอ) · บันทึก Observe (O ใน POE) ลงใบงานพร้อมตารางข้อมูล/กราฟตามที่กำหนด",
        "goal": "สนับสนุนทักษะกระบวนการวิทยาศาสตร์และการเรียนรู้ผ่านประสบการณ์ตรง (Experience Before Labeling) · นักเรียนสร้างความรู้เอง · สอดคล้อง วPA ตัวชี้วัด ๓ (สร้างความรู้เอง) และ ๗ (จัดการตนเอง)",
    },
    "E3": {
        "icon": "💬",
        "name": "Explain · อธิบายและสังเคราะห์",
        "teacher": "ครูเปิดเวที CER Live Board ให้นักเรียนแต่ละกลุ่มนำเสนอ Claim-Evidence-Reasoning ที่ได้จากขั้น Explore · ครูเชื่อมโยงข้อมูลของแต่ละกลุ่มเข้าสู่หลักการทางวิทยาศาสตร์ที่ถูกต้อง และแก้ misconception เป้าหมายอย่างชัดเจน",
        "student": "นักเรียนรายกลุ่มนำเสนอ Claim/Evidence/Reasoning บนกระดาน Live Board หรือ Padlet · นักเรียนรายบุคคลเขียน Explain (E ใน POE) เปรียบเทียบกับ Predict ตอนต้น · ตั้งคำถามต่อกลุ่มอื่น",
        "goal": "นักเรียนสร้างคำอธิบายเชิงวิทยาศาสตร์จากหลักฐานของตน · ครูฝัง concept ทางวิทยาศาสตร์อย่างเป็นทางการในจังหวะที่นักเรียนพร้อม · สอดคล้อง วPA ตัวชี้วัด ๓ (สร้างความรู้เอง) · ๕ (พัฒนาทักษะนวัตกรรม) · ๖ (ข้อมูลย้อนกลับ)",
    },
    "E4": {
        "icon": "🚀",
        "name": "Elaborate · ขยายผลและประยุกต์",
        "teacher": "ครูแจกใบงานคำนวณ (Calc) ไล่ระดับความยาก และใบงาน Spot the Error ที่นักเรียนต้องวิเคราะห์ misconception ในสถานการณ์สมมุติ · ครูเปิดโอกาสให้นักเรียนเลือกระดับโจทย์ (differentiation)",
        "student": "นักเรียนรายบุคคลทำใบงาน Calc และ Spot the Error ตามลำดับ · ปรึกษาเพื่อนคู่ข้างได้สำหรับข้อ Challenge · บันทึกวิธีคิดและเหตุผลของแต่ละข้อ ไม่ใช่แค่คำตอบ",
        "goal": "นักเรียนนำความรู้ไปประยุกต์ใช้ในบริบทใหม่ (Far Transfer) · ฝึกทักษะการแก้โจทย์ในสถานการณ์ PISA และการสอบเข้ามหาวิทยาลัย · สอดคล้อง วPA ตัวชี้วัด ๕ (พัฒนาทักษะนวัตกรรม) และ ๘ (นำไปต่อยอด)",
    },
    "F-Post": {
        "icon": "📊",
        "name": "Formative Post-test · Four-tier Diagnostic",
        "teacher": "ครูเปิดข้อสอบ Post-test (parallel form ของ Pre-test) บนระบบ Apps Script · ครูดู dashboard ภาพรวมและระบุ misconception ที่ลดมากที่สุดและน้อยที่สุด เพื่อปรับการสอนครั้งถัดไป",
        "student": "นักเรียนรายบุคคลทำ Post-test 5 ข้อบนอุปกรณ์ · ครบ 4 ระดับเหมือน Pre-test · ระบบคำนวณ Hake gain อัตโนมัติเทียบกับคะแนน Pre-test ของตน",
        "goal": "วัด Conceptual Change ของนักเรียน · ครูได้หลักฐานเชิงปริมาณของการเปลี่ยนแปลงเชิงมโนทัศน์ (% Sound เพิ่ม · % Misconception ลด) · สอดคล้อง วPA ตัวชี้วัด ๖ (ข้อมูลย้อนกลับ)",
    },
    "MJ": {
        "icon": "📔",
        "name": "Metacognitive Journal 3-2-1 · สะท้อนคิด",
        "teacher": "ครูเปิดใบบันทึก Metacognitive Journal · กระตุ้นนักเรียนสะท้อนคิดเชิงลึก ไม่ใช่แค่สรุปเนื้อหา · ครูใช้ข้อมูลสะท้อนคิดนี้ปรับการสอนครั้งถัดไป (Responsive Teaching)",
        "student": "นักเรียนรายบุคคลเขียน MJ-3-2-1 — <strong>3:</strong> สิ่งที่เข้าใจวันนี้ (ระบุชัดเจน) · <strong>2:</strong> ปรากฏการณ์ในชีวิตจริงที่เชื่อมโยงได้ · <strong>1:</strong> คำถามค้างใจที่อยากรู้ต่อ · ชู Traffic Light Post อีกครั้ง",
        "goal": "นักเรียนพัฒนา metacognition · ครูได้ข้อมูลย้อนกลับเชิงคุณภาพและคำถาม hook สู่แผนถัดไป · สอดคล้อง วPA ตัวชี้วัด ๖ (ข้อมูลย้อนกลับ) และ ๘ (นำไปต่อยอด)",
    },
}

def get_phase_info(code):
    code_u = code.upper().strip()
    if code_u.startswith("F-PRE") or code_u == "F-PRE": return PHASE_INFO["F-Pre"]
    if code_u.startswith("F-POST") or code_u in ("F-POST","FT-02","F-FINAL"): return PHASE_INFO["F-Post"]
    if code_u in PHASE_INFO: return PHASE_INFO[code_u]
    # generic
    return {
        "icon": "▌",
        "name": code,
        "teacher": "ครูดำเนินกิจกรรมตามรายละเอียดต่อไปนี้",
        "student": "นักเรียนปฏิบัติกิจกรรมตามที่ครูกำหนด",
        "goal": "เพื่อให้บรรลุจุดประสงค์การเรียนรู้ของขั้นนี้",
    }

# ─────────────────────────────────────────────────────────────
# Build one DOCX (เวอร์ชันใหม่ · บรรยายความ + วPA)
def build_docx(unit_info, html_path, out_path):
    unit_key, unit_full, chapter_text, grade, indicator, _ = unit_info
    data = parse_infographic(html_path)
    doc = Document()

    # margins
    for sec in doc.sections:
        sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2.0); sec.right_margin = Cm(1.8)

    plan_num_th = thai_num(data.get('plan_num','-'))

    # ────────────── COVER (2-column table)
    cover_t = doc.add_table(rows=1, cols=2)
    cover_t.autofit = False
    cover_t.columns[0].width = Cm(4.5)
    cover_t.columns[1].width = Cm(12.5)
    left = cover_t.rows[0].cells[0]
    right = cover_t.rows[0].cells[1]
    left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(left, "1f4e79")
    # left: plan number
    left.text = ""
    p = left.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("แผนการจัดการเรียนรู้\nที่ "); set_th_font(r, size=14, bold=True, color=(255,255,255))
    p2 = left.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run(plan_num_th); set_th_font(r, size=42, bold=True, color=(255,255,255))
    # right: title + info
    right.text = ""
    p = right.paragraphs[0]
    r = p.add_run("รายวิชา ว 30203 ฟิสิกส์ 3"); set_th_font(r, size=13, bold=True)
    p = right.add_paragraph()
    r = p.add_run(f"ระดับชั้น {grade}   ภาคเรียนที่ ๑   ปีการศึกษา ๒๕๖๙"); set_th_font(r, size=13)
    p = right.add_paragraph()
    r = p.add_run(f"{unit_full}   {chapter_text}"); set_th_font(r, size=13)
    p = right.add_paragraph()
    r = p.add_run("เรื่อง "); set_th_font(r, size=14, bold=True)
    r = p.add_run(data.get("title","-")); set_th_font(r, size=16, bold=True, color=(0x1f,0x4e,0x79))
    p = right.add_paragraph()
    if data.get("subtitle"):
        r = p.add_run(data["subtitle"]); set_th_font(r, size=12, color=(0x55,0x55,0x55), italic=True)
    p = right.add_paragraph()
    period_line = "  ·  ".join(data.get("meta", [])[:3])
    r = p.add_run(period_line); set_th_font(r, size=12)
    p = right.add_paragraph()
    r = p.add_run("ครูผู้สอน นายโกเมน ปาปะโถ   กลุ่มสาระการเรียนรู้วิทยาศาสตร์และเทคโนโลยี"); set_th_font(r, size=12)
    p = right.add_paragraph()
    r = p.add_run("โรงเรียนสตรีวิทยา"); set_th_function = set_th_font(r, size=12, bold=True)

    add_para(doc, "", size=10, space_after=4)

    # ────────────── ๑. สาระสำคัญ
    add_h1(doc, "๑. สาระสำคัญ")
    concept = data.get("concept", "")
    lines = [l.strip(" •·-") for l in concept.split("\n") if l.strip()]
    if lines:
        add_para(doc, lines[0], size=SZ_BODY, space_after=4,
                 indent_left=IND_BODY, first_line_indent=IND_BODY_FL, line_spacing=1.5)
        for line in lines[1:]:
            add_bullet(doc, line)

    # ────────────── ๒. มาตรฐานการเรียนรู้และตัวชี้วัด
    add_h1(doc, "๒. มาตรฐานการเรียนรู้และตัวชี้วัด")
    add_para(doc, "กลุ่มสาระการเรียนรู้วิทยาศาสตร์และเทคโนโลยี · สาระที่ 5 พลังงาน (สำหรับหน่วยคลื่น/SHM/เสียง) หรือสาระที่ 2 วิทยาศาสตร์กายภาพ (สำหรับหน่วยแสง) ตามหลักสูตรแกนกลางการศึกษาขั้นพื้นฐาน พุทธศักราช 2551 (ฉบับปรับปรุง พ.ศ. 2560)",
             size=SZ_BODY, indent_left=IND_BODY, first_line_indent=IND_BODY_FL, line_spacing=1.5)
    add_h3(doc, "ตัวชี้วัด/ผลการเรียนรู้")
    add_bullet(doc, indicator)
    add_bullet(doc, f"อ้างอิงเนื้อหา: {chapter_text} (สสวท. ฟิสิกส์ เล่ม 3 · ม.5)")

    # ────────────── ๓. จุดประสงค์การเรียนรู้ (K/P/A)
    add_h1(doc, "๓. จุดประสงค์การเรียนรู้ (K · P · A)")
    add_para(doc, "จุดประสงค์การเรียนรู้แบ่งเป็น 3 ด้านตามแนวทาง KPA โดยแต่ละจุดประสงค์มีรหัสและเกณฑ์การวัดที่ชัดเจน เพื่อให้สามารถประเมินได้อย่างเป็นรูปธรรม",
             size=SZ_BODY, indent_left=IND_BODY, first_line_indent=IND_BODY_FL, line_spacing=1.5)
    rows = []
    for k in data["K"]: rows.append(["K (ความรู้)", k["code"], k["text"], "Pre/Post-test ≥ 80% Sound"])
    for p_ in data["P"]: rows.append(["P (ทักษะ)", p_["code"], p_["text"], "Rubric ≥ 2.0/3.0"])
    for a in data["A"]: rows.append(["A (เจตคติ)", a["code"], a["text"], "OB/MJ ≥ 70%"])
    if rows:
        add_table(doc, ["ด้าน", "รหัส", "จุดประสงค์การเรียนรู้", "เกณฑ์ผ่าน/เครื่องมือ"],
                  rows, col_widths_cm=[2.2, 1.5, 8.5, 4.5])

    # ────────────── ๔. สมรรถนะของผู้เรียน
    add_h1(doc, "๔. สมรรถนะของผู้เรียน")
    add_para(doc, "แผนนี้ออกแบบให้นักเรียนได้พัฒนาสมรรถนะสำคัญตามหลักสูตรแกนกลาง 6 ด้าน ดังนี้",
             size=SZ_BODY, indent_left=IND_BODY, first_line_indent=IND_BODY_FL, line_spacing=1.5)
    for i, c in enumerate([
        "การสื่อสาร — นำเสนอ Claim/Evidence/Reasoning ในขั้น Explain ได้ชัดเจน รับ-ส่งสารกับเพื่อนและครูอย่างมีประสิทธิภาพ",
        "การคิด — วิเคราะห์ผลการทดลอง เชื่อมโยงข้อมูลจาก POE และตัดสินใจเลือกแนวทางแก้ปัญหา",
        "การแก้ปัญหา — แก้โจทย์เชิงคำนวณและสถานการณ์ Spot the Error ในระดับความยากต่าง ๆ",
        "ทักษะชีวิตและการใช้เทคโนโลยี — ใช้ Simulation และเครื่องมือ Lab อย่างปลอดภัยและคุ้มค่า",
        "การรวมพลังทำงานเป็นทีม — แบ่งหน้าที่และจัดการเวลาในกลุ่มทดลอง 3 ฐานอย่างเป็นระบบ",
    ], 1):
        add_numbered(doc, thai_num(i), c)

    # ────────────── ๕. คุณลักษณะอันพึงประสงค์
    add_h1(doc, "๕. คุณลักษณะอันพึงประสงค์")
    add_para(doc, "ครูเฝ้าสังเกตและบันทึกคุณลักษณะของนักเรียนตลอดคาบเรียนด้วย OB sheet โดยเฉพาะ:",
             size=SZ_BODY, indent_left=IND_BODY, first_line_indent=IND_BODY_FL, line_spacing=1.5)
    for c in ["มีวินัย — เข้าร่วมกิจกรรมตรงเวลา · ทำงานในเวลาที่กำหนด",
              "ใฝ่เรียนรู้ — ตั้งคำถาม · ค้นคว้าเพิ่มเติม · ทำการบ้านครบ",
              "มุ่งมั่นในการทำงาน — ไม่ท้อกับโจทย์ยาก · ลองหลายวิธี",
              "ซื่อสัตย์ทางวิทยาศาสตร์ — รายงานผลการทดลองตามที่วัดได้จริง · ยอมรับข้อมูลที่ขัดกับความเชื่อเดิม"]:
        add_bullet(doc, c)

    # ────────────── ๖. สาระการเรียนรู้และ Misconceptions เป้าหมาย
    add_h1(doc, "๖. สาระการเรียนรู้และ Misconceptions เป้าหมาย")
    add_h3(doc, "๖.๑ เนื้อหา")
    if data.get("refr_cards"):
        for i, c in enumerate(data["refr_cards"], 1):
            txt = f"({thai_num(i)}) {c['title']} — {c['sub']}"
            if c["note"]: txt += f" · {c['note']}"
            add_para(doc, txt, size=SZ_BODY, space_after=3,
                     indent_left=IND_BODY, line_spacing=1.5)
    if data.get("formulas"):
        add_h3(doc, "สมการสำคัญ")
        for f in data["formulas"]:
            line = f"[{f['tag']}]  {f['formula']}"
            if f["note"]: line += f"   ({f['note']})"
            add_bullet(doc, line)

    if data["misc"]:
        add_h3(doc, "๖.๒ Misconceptions เป้าหมาย")
        add_para(doc, "ครูออกแบบคาบนี้เพื่อแก้ misconception ที่ฝังลึกของนักเรียน โดยใช้ Concept Cartoon ในขั้น Engage และ Four-tier Diagnostic ในขั้น Pre/Post เพื่อวัดการเปลี่ยนแปลงเชิงมโนทัศน์ (Conceptual Change) อย่างเป็นรูปธรรม",
                 size=SZ_BODY, indent_left=IND_BODY, first_line_indent=IND_BODY_FL, space_after=6, line_spacing=1.5)
        mrows = []
        for m in data["misc"]:
            mrows.append([m["id"], "★" if "★" in m["id"] or "1" in m["id"] else "", m["wrong"], m["correct"]])
        add_table(doc, ["รหัส", "ระดับ", "นักเรียนมักเข้าใจว่า...", "แนวคิดที่ถูกต้อง"],
                  mrows, col_widths_cm=[1.8, 1.2, 6.5, 7.2])

    # ────────────── ๗. กิจกรรมการเรียนรู้
    add_pagebreak(doc)
    add_h1(doc, "๗. กิจกรรมการเรียนรู้ (5E + POE)")
    add_para(doc, "รูปแบบการจัดการเรียนรู้: สืบเสาะหาความรู้ 5E ของ BSCS (Engage → Explore → Explain → Elaborate → Evaluate) ร่วมกับเทคนิค Predict–Observe–Explain (POE) ของ White & Gunstone (1992) รวม 9 ขั้นย่อย ใช้เวลาทั้งหมด 100 นาที (2 คาบ ๆ ละ 50 นาที)",
             size=SZ_BODY, indent_left=IND_BODY, first_line_indent=IND_BODY_FL, space_after=4, line_spacing=1.5)
    add_para(doc, "ภาษาที่ใช้: Child-centered บรรยายตาม 5W1H (ใคร ทำอะไร ที่ไหน อย่างไร เพื่ออะไร ภายในเวลาเท่าใด) แต่ละย่อหน้าระบุชัด — ครูทำอะไร · นักเรียนทำอะไร · เพื่อบรรลุเป้าหมายใด",
             size=SZ_NOTE+1, indent_left=IND_BODY, space_after=8, line_spacing=1.4,
             italic=True, color=(0x55,0x55,0x55))

    for idx, ph in enumerate(data["phases"], 1):
        info = get_phase_info(ph["code"])
        # หัวข้อกิจกรรม
        title = f"กิจกรรม ๗.{thai_num(idx)} · {info['icon']} ขั้น {ph['code']} — {info['name']}"
        if ph.get("time"): title += f"   ⏱ {ph['time']}"
        add_h2(doc, title)

        # tag line
        tags = []
        if ph.get("obj"): tags.append("จุดประสงค์: " + ", ".join(ph["obj"]))
        if ph.get("misc"): tags.append("Misconception: " + ", ".join(ph["misc"]))
        if ph.get("vpa"): tags.append("วPA: " + ", ".join(ph["vpa"]))
        if ph.get("poe"): tags.append("POE: " + ph["poe"])
        if tags:
            add_para(doc, "   ".join("[" + t + "]" for t in tags),
                     size=SZ_TAGS, color=(0x55,0x55,0x55), italic=True,
                     indent_left=IND_BODY, space_after=6)

        # บรรยาย 3 ย่อหน้า · ครู / นักเรียน / เป้าหมาย
        for label, txt, color in [
            ("ครู", info["teacher"], (0x1f,0x4e,0x79)),
            ("นักเรียน", info["student"], (0x2e,0x7d,0x32)),
            ("เป้าหมาย", info["goal"], (0xb7,0x4d,0x00)),
        ]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Cm(IND_BODY)
            p.paragraph_format.line_spacing = 1.5
            r = p.add_run(f"{label}: "); set_th_font(r, size=SZ_BODY, bold=True, color=color)
            r = p.add_run(th_break(strip_html(txt))); set_th_font(r, size=SZ_BODY)

        # 4) รายละเอียดกิจกรรมเฉพาะแผนนี้
        if ph.get("activities"):
            add_h3(doc, "ลำดับกิจกรรมในขั้นนี้", indent=IND_BODY)
            for i, act in enumerate(ph["activities"], 1):
                add_numbered(doc, thai_num(i), act)

    # ────────────── ๘. สื่อ อุปกรณ์ และแหล่งการเรียนรู้
    add_pagebreak(doc)
    add_h1(doc, "๘. สื่อ อุปกรณ์ และแหล่งการเรียนรู้")
    add_para(doc, "สื่อและเครื่องมือต่อไปนี้ออกแบบให้ทำงานสอดคล้องกันตลอดคาบ — Concept Cartoon ในขั้น Engage · ใบบันทึก POE ในขั้น Explore-Explain · ใบงาน Calc และ Spot the Error ในขั้น Elaborate · Four-tier Diagnostic ก่อน-หลังเรียน · TL-card และ MJ-journal สำหรับ formative assessment",
             size=SZ_BODY, indent_left=IND_BODY, first_line_indent=IND_BODY_FL, space_after=6, line_spacing=1.5)
    if data.get("sim"):
        add_h3(doc, "Simulation หลัก")
        add_bullet(doc, f"{data['sim']['title']}")
        add_bullet(doc, f"URL: {data['sim']['url']}", size=SZ_TABLE)
        if data['sim'].get("badges"):
            add_bullet(doc, "โหมดที่ใช้: " + " · ".join(data["sim"]["badges"]), size=SZ_TABLE)
    add_h3(doc, "รายการสื่อและอุปกรณ์")
    for i, tool in enumerate(data.get("tools", []), 1):
        add_numbered(doc, thai_num(i), tool)

    # ────────────── ๙. การวัดและประเมินผล
    add_h1(doc, "๙. การวัดและประเมินผล")
    add_para(doc, "แผนนี้ใช้ทั้ง Formative Assessment (TL-card · Four-tier Pre/Post · POE Rubric · MJ Journal) และ Summative Assessment (ใบงาน Calc + Spot the Error) เพื่อตอบโจทย์การประเมิน 3 ด้าน — การเปลี่ยนแปลงเชิงมโนทัศน์ · ทักษะกระบวนการ · เจตคติ — อย่างครบถ้วน ครูสามารถดูผลประเมินอัตโนมัติได้ที่ KP-Classroom Teacher Dashboard ผ่านระบบ Apps Script",
             size=SZ_BODY, indent_left=IND_BODY, first_line_indent=IND_BODY_FL, space_after=6, line_spacing=1.5)
    if data.get("assess"):
        add_table(doc, ["สิ่งที่ต้องการวัด", "ตัวแปร", "เครื่องมือ", "เกณฑ์ผ่าน"],
                  data["assess"], col_widths_cm=[5.0, 4.0, 4.0, 3.7])
    if data.get("rubric"):
        add_h3(doc, "เกณฑ์ Rubric ของกิจกรรมหลัก (0–3 คะแนน)")
        for r in data["rubric"]:
            add_bullet(doc, f"คะแนน {r['score']} — {r['desc']}")

    # ────────────── ๑๐. การเชื่อมโยงกับ วPA ด้านที่ 1
    add_pagebreak(doc)
    add_h1(doc, "๑๐. การเชื่อมโยงกับ วPA ด้านที่ 1")
    add_para(doc, "แผนนี้ได้รับการออกแบบให้ตอบโจทย์ตัวชี้วัด วPA ด้านที่ 1 ทักษะการจัดการเรียนรู้และการจัดการชั้นเรียน ครบทั้ง 8 ตัวชี้วัด หลักฐานทั้งหมดสามารถนำไปประกอบรายงานการประเมินวิทยฐานะเชี่ยวชาญได้โดยตรง โดยเชื่อมโยงกับกิจกรรมและสื่อในแต่ละขั้นของ 5E ดังตารางต่อไปนี้",
             size=SZ_BODY, indent_left=IND_BODY, first_line_indent=IND_BODY_FL, space_after=6, line_spacing=1.5)
    if data.get("vpa"):
        rows = []
        for v in data["vpa"]:
            if len(v) >= 4: rows.append([v[0], v[1], v[2], v[3]])
            elif len(v) >= 3: rows.append([v[0], v[1], v[2], "—"])
        add_table(doc, ["#", "ตัวชี้วัด วPA ด้านที่ 1", "หลักฐาน/กิจกรรม", "ขั้น 5E"],
                  rows, col_widths_cm=[1.0, 5.5, 6.5, 3.7])

    # ────────────── ๑๑. บันทึกหลังการสอน
    add_pagebreak(doc)
    add_h1(doc, "๑๑. บันทึกหลังการจัดการเรียนรู้")
    add_para(doc, "ครูบันทึกผลการสอนทันทีหลังจบคาบ เพื่อสะท้อนคิดและปรับการสอนครั้งต่อไป (Reflective Practice) ข้อมูลในส่วนนี้เป็นหลักฐานสำคัญสำหรับงานวิจัย วPA ของครูเอง",
             size=SZ_BODY, indent_left=IND_BODY, first_line_indent=IND_BODY_FL,
             space_after=6, italic=True, color=(0x55,0x55,0x55), line_spacing=1.5)
    for label in ["๑๑.๑ ผลที่เกิดกับผู้เรียน — ด้านความรู้ (K)",
                  "๑๑.๒ ผลที่เกิดกับผู้เรียน — ด้านทักษะกระบวนการ (P)",
                  "๑๑.๓ ผลที่เกิดกับผู้เรียน — ด้านเจตคติ (A)",
                  "๑๑.๔ ปัญหา/อุปสรรค ระหว่างการสอน",
                  "๑๑.๕ แนวทางปรับปรุง/พัฒนา ในการสอนครั้งต่อไป",
                  "๑๑.๖ ข้อเสนอแนะของผู้บริหาร/ผู้นิเทศ"]:
        add_h3(doc, label, indent=IND_BODY)
        for _ in range(3):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(IND_BODY)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run("................................................................................................................................")
            set_th_font(r, size=SZ_BODY, color=(0x99,0x99,0x99))

    # ลงนาม
    add_para(doc, "", size=14, space_before=12)
    sig_t = doc.add_table(rows=1, cols=2)
    sig_t.autofit = False
    sig_t.columns[0].width = Cm(8.5)
    sig_t.columns[1].width = Cm(8.5)
    for cell, role, name in [(sig_t.rows[0].cells[0], "ผู้สอน", "นายโกเมน ปาปะโถ"),
                              (sig_t.rows[0].cells[1], "ผู้นิเทศ", "............................................")]:
        cell.text = ""
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("ลงชื่อ ................................................"); set_th_font(r, size=13)
        p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"({name})"); set_th_font(r, size=13)
        p = cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(role); set_th_font(r, size=13, bold=True)

    doc.save(str(out_path))
    return out_path

# ─────────────────────────────────────────────────────────────
def find_infographic(unit_key, plan_num):
    unit_dir = BASE / unit_key
    for pdir in sorted(unit_dir.iterdir()):
        if not pdir.is_dir(): continue
        name = pdir.name
        if name.startswith(f"แผน{plan_num:02d}") or name.startswith(f"แผน{plan_num}"):
            for f in pdir.glob("*Infographic*.html"):
                return f
            for f in pdir.glob("*infographic*.html"):
                return f
    return None

def main():
    OUTDIR.mkdir(exist_ok=True)
    summary = []
    for unit in UNITS:
        key, _, _, _, _, n = unit
        for i in range(1, n+1):
            html = find_infographic(key, i)
            if not html:
                summary.append(f"  ❌ {key} แผน{i}: ไม่พบ infographic")
                continue
            outname = f"{key}_แผน{i:02d}_{html.parent.name.split('_',1)[-1]}.docx"
            outpath = OUTDIR / outname
            try:
                build_docx(unit, html, outpath)
                summary.append(f"  ✅ {outname}")
            except Exception as e:
                summary.append(f"  ❌ {key} แผน{i}: {e}")
                import traceback; traceback.print_exc()
    print("\n".join(summary))
    print(f"\nรวม {len([s for s in summary if s.startswith('  ✅')])} ไฟล์ใน {OUTDIR}")

if __name__ == "__main__":
    main()
